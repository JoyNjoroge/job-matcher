"""
Parser Service - CV/resume extraction with optional OpenRouter enrichment.
"""

import io
import re
import json
from PyPDF2 import PdfReader
from docx import Document
import os

def parse_cv(file) -> str:
    """
    Parse CV file and extract text content.
    
    Args:
        file: File object (PDF or DOCX)
    
    Returns:
        Extracted text content as string
    """
    filename = file.filename.lower()
    
    try:
        if filename.endswith(".pdf"):
            return parse_pdf(file)
        elif filename.endswith(".docx"):
            return parse_docx(file)
        else:
            # Try to read as plain text
            content = file.read()
            if isinstance(content, bytes):
                return content.decode("utf-8")
            return content
    except Exception as e:
        print(f"CV parsing error: {e}")
        return ""


def parse_resume_file(file, file_ext: str) -> str:
    """
    Parse resume file by extension.
    
    Args:
        file: File object
        file_ext: File extension (pdf, docx, etc.)
    
    Returns:
        Extracted text content
    """
    try:
        if file_ext == "pdf":
            return parse_pdf(file)
        elif file_ext in ["docx", "doc"]:
            return parse_docx(file)
        else:
            content = file.read()
            if isinstance(content, bytes): 
                return content.decode("utf-8", errors='ignore')
            return content
    except Exception as e:
        print(f"Resume parsing error: {e}")
        return ""


def parse_pdf(file) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file: PDF file object
    
    Returns:
        Extracted text content
    """
    try:
        # Reset file pointer to beginning
        file.seek(0)
        reader = PdfReader(io.BytesIO(file.read()))
        text_content = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
        
        return "\n".join(text_content)
    except Exception as e:
        print(f"PDF parsing error: {e}")
        return ""


def parse_docx(file) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file: DOCX file object
    
    Returns:
        Extracted text content
    """
    try:
        # Reset file pointer to beginning
        file.seek(0)
        doc = Document(io.BytesIO(file.read()))
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    except Exception as e:
        print(f"DOCX parsing error: {e}")
        return ""


def extract_resume_structure(raw_text: str) -> dict:
    """
    Extract structured data from raw resume text using pattern matching.
    This is the deterministic fallback when an AI provider is unavailable.
    
    Args:
        raw_text: Raw text content from resume
    
    Returns:
        Structured dictionary with resume sections
    """
    structure = {
        "education": [],
        "experience": [],
        "skills": [],
        "projects": [],
        "tools": [],
        "certifications": [],
        "languages": [],
        "awards": [],
        "volunteer_experience": [],
        "publications": [],
        "courses": [],
        "interests": [],
        "additional_details": {},
        "years_of_experience": None,
        "experience_level": None,
        "summary": None,
        "full_name": None,
        "phone": None,
        "email": None,
        "location": None,
        "job_titles": [],
    }
    
    if not raw_text:
        return structure
    
    lines = raw_text.split("\n")
    
    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, raw_text)
    if emails:
        structure["email"] = emails[0]
    
    # Extract phone
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, raw_text)
    if phones:
        structure["phone"] = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
    
    # Try to extract name (usually first line or near top)
    first_lines = [line.strip() for line in lines[:5] if line.strip()]
    if first_lines:
        # Name is usually short, title-cased, and near the top
        for line in first_lines:
            if len(line.split()) <= 4 and line[0].isupper() and not any(c.isdigit() for c in line):
                structure["full_name"] = line
                break
    
    # Extract skills using common patterns
    skills_patterns = [
        r'\b(?:python|javascript|typescript|java|c\+\+|c#|ruby|go|rust|swift|kotlin|php|scala)\b',
        r'\b(?:react|angular|vue|node\.?js|express|django|flask|spring|laravel)\b',
        r'\b(?:aws|azure|gcp|docker|kubernetes|terraform|jenkins)\b',
        r'\b(?:sql|postgresql|mysql|mongodb|redis|elasticsearch)\b',
        r'\b(?:git|ci/cd|agile|scrum|jira)\b',
        r'\b(?:html|css|sass|tailwind|bootstrap)\b',
        r'\b(?:rest|graphql|api|microservices)\b',
    ]
    
    extracted_skills = set()
    for pattern in skills_patterns:
        matches = re.findall(pattern, raw_text.lower())
        extracted_skills.update(matches)
    
    structure["skills"] = list(extracted_skills)
    
    # Estimate years of experience from date ranges
    year_patterns = re.findall(r'(\d{4})\s*[-–—]\s*(\d{4}|present|current)', raw_text.lower())
    if year_patterns:
        total_years = 0
        current_year = 2026
        for start, end in year_patterns:
            start_year = int(start)
            end_year = current_year if end in ["present", "current"] else int(end)
            total_years += max(0, end_year - start_year)
        
        structure["years_of_experience"] = min(total_years, 50)  # Cap at reasonable max
        
        # Estimate experience level
        if total_years < 2:
            structure["experience_level"] = "entry"
        elif total_years < 5:
            structure["experience_level"] = "mid"
        elif total_years < 10:
            structure["experience_level"] = "senior"
        else:
            structure["experience_level"] = "lead"
    
    # Extract job titles (look for lines that might be job titles)
    title_keywords = ['engineer', 'developer', 'manager', 'analyst', 'designer', 'consultant', 
                     'architect', 'lead', 'senior', 'junior', 'director', 'specialist']
    
    for line in lines:
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in title_keywords):
            # Clean and add as potential job title
            cleaned = line.strip()
            if len(cleaned) < 80 and len(cleaned.split()) <= 8:
                structure["job_titles"].append(cleaned)
    
    structure["job_titles"] = list(set(structure["job_titles"]))[:5]  # Limit to 5 unique titles
    
    # Extract first meaningful paragraph as summary
    paragraphs = [p.strip() for p in raw_text.split('\n\n') if len(p.strip()) > 100]
    if paragraphs:
        # Find paragraph that looks like a professional summary
        for para in paragraphs[:3]:
            if len(para) > 100 and not any(
                keyword in para.lower() 
                for keyword in ["education", "university", "degree", "bachelor", "master"]
            ):
                structure["summary"] = para[:500]  # First 500 chars
                break
    
    return structure


def use_ai_for_parsing(raw_text: str) -> dict:
    """
    Use the configured AI provider for richer parsing, with a local fallback.
    
    Args:
        raw_text: Raw text content from resume
    
    Returns:
        Structured dictionary with resume sections
    """
    try:
        if not os.getenv("OPENROUTER_API_KEY"):
            print("No OpenRouter key, using fallback parser")
            return extract_resume_structure(raw_text)
        from services.ai import _generate_content_text
        
        prompt = f"""
        Parse this resume comprehensively. Preserve every factual detail that
        could help complete a job application. Do not invent or infer facts
        that are not supported by the resume.
        Return ONLY valid JSON (no markdown, no explanation) with this exact structure:
        {{
            "full_name": "string or null",
            "email": "string or null",
            "phone": "string or null",
            "location": "string or null",
            "summary": "string or null",
            "job_titles": ["current/recent job titles"],
            "skills": ["skill1", "skill2"],
            "tools": ["tool1", "tool2"],
            "education": [
                {{"institution": "string", "degree": "string", "field": "string", "start_date": "string", "end_date": "string", "grade": "string", "location": "string", "details": ["string"]}}
            ],
            "work_experience": [
                {{"company": "string", "title": "string", "location": "string", "start_date": "string", "end_date": "string", "description": "string", "achievements": ["string"], "technologies": ["string"]}}
            ],
            "projects": [
                {{"name": "string", "role": "string", "url": "string", "start_date": "string", "end_date": "string", "description": "string", "highlights": ["string"], "technologies": ["string"]}}
            ],
            "certifications": [
                {{"name": "string", "issuer": "string", "date": "string", "expiry_date": "string", "credential_id": "string", "url": "string"}}
            ],
            "languages": [{{"name": "string", "proficiency": "string"}}],
            "awards": [{{"name": "string", "issuer": "string", "date": "string", "description": "string"}}],
            "volunteer_experience": [{{"organization": "string", "role": "string", "start_date": "string", "end_date": "string", "description": "string"}}],
            "publications": [{{"title": "string", "publisher": "string", "date": "string", "url": "string", "description": "string"}}],
            "courses": [{{"name": "string", "provider": "string", "date": "string"}}],
            "interests": ["string"],
            "years_of_experience": 0,
            "experience_level": "entry|mid|senior|lead|executive",
            "additional_details": {{
                "key copied from an otherwise unclassified resume section": "complete factual value"
            }}
        }}
        
        Resume text:
        {raw_text[:16000]}
        """
        
        response_text = _generate_content_text(
            prompt,
            json_mode=True,
            max_tokens=4096,
        )
        
        # Remove markdown code blocks if present
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        response_text = response_text.strip()
        
        # Parse the JSON
        result = json.loads(response_text)
        
        # Ensure all required fields exist
        defaults = {
            "full_name": None,
            "email": None,
            "phone": None,
            "location": None,
            "summary": None,
            "job_titles": [],
            "skills": [],
            "tools": [],
            "education": [],
            "work_experience": [],
            "projects": [],
            "certifications": [],
            "languages": [],
            "awards": [],
            "volunteer_experience": [],
            "publications": [],
            "courses": [],
            "interests": [],
            "additional_details": {},
            "years_of_experience": None,
            "experience_level": None,
        }
        
        for key, default_value in defaults.items():
            if key not in result:
                result[key] = default_value

        # Accept older model output while keeping one canonical profile field.
        if not result.get("work_experience") and result.get("experience"):
            result["work_experience"] = result["experience"]
        result.pop("experience", None)
        
        print(f"AI parsing successful: extracted {len(result.get('skills', []))} skills")
        return result
        
    except json.JSONDecodeError as e:
        print(f"AI JSON decode error: {e}")
        print(f"Response was: {response_text[:200]}...")
        return extract_resume_structure(raw_text)
    except Exception as e:
        print(f"AI parsing error: {e}")
        import traceback
        traceback.print_exc()
        return extract_resume_structure(raw_text)
